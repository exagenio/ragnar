import json
import base64
import struct
from io import BytesIO
from typing import Callable

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import plotly.io as pio


def generate_report_document(report, sections, progress_callback: Callable[[str], None] | None = None):
    """Generate report document"""
    doc = Document()
    progress = progress_callback or (lambda message: None)

    progress("Building title page.")
    _add_title_page(doc, report)
    doc.add_page_break()

    progress("Building table of contents.")
    _add_table_of_contents(doc, sections)
    doc.add_page_break()

    total_sections = len(sections)
    for section_num, section in enumerate(sections, start=1):
        progress(f"Writing section {section_num} of {total_sections}: {section.title}")
        _add_section(doc, section, report, section_num, progress)

    progress("Finalizing DOCX package.")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


def _add_title_page(doc, report):
    """Add title page"""
    title = doc.add_heading(report.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = title.runs[0]
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)


def _add_table_of_contents(doc, sections):
    """Add table of contents"""

    heading = doc.add_heading("Table of Contents", level=1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    for section_num, section in enumerate(sections, start=1):
        section_para = doc.add_paragraph()
        run = section_para.add_run(f"• {section.title}")
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

        for subsection_num, subsection in enumerate(section.sub_sections.all(), start=1):
            sub_para = doc.add_paragraph()
            sub_para.paragraph_format.left_indent = Inches(0.3)
            run = sub_para.add_run(f"◦ {subsection.title}")
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(51, 102, 153)

            for topic_num, topic in enumerate(subsection.topics.filter(is_approved=True), start=1):
                topic_para = doc.add_paragraph()
                topic_para.paragraph_format.left_indent = Inches(0.6)
                topic_number = f"{section_num}.{subsection_num}.{topic_num}"
                run = topic_para.add_run(f"{topic_number} {topic.title}")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(102, 153, 204)

        doc.add_paragraph()


def _add_section(doc, section, report, section_num, progress_callback):
    """Add section"""

    heading = doc.add_heading(f"{section_num}. {section.title}", level=1)
    run = heading.runs[0]

    style_heading(run, 1)
    run.font.color.rgb = RGBColor(0, 51, 102)

    if hasattr(section, 'content') and section.content.status == 'generated':
        _add_section_content(doc, section.content.content_json)

    for subsection_num, subsection in enumerate(section.sub_sections.all(), start=1):
        progress_callback(
            f"Writing subsection {section_num}.{subsection_num}: {subsection.title}"
        )
        _add_subsection(doc, subsection, report, section_num, subsection_num, progress_callback)

    doc.add_page_break()


def _add_section_content(doc, content_json):
    """Add section content"""

    if 'section_introduction' in content_json:
        for text in content_json['section_introduction'].get('paragraphs', []):
            para = doc.add_paragraph(text)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _add_subsection(doc, subsection, report, section_num, subsection_num, progress_callback):
    """Add subsection"""

    heading = doc.add_heading(f"{section_num}.{subsection_num} {subsection.title}", level=2)
    run = heading.runs[0]

    style_heading(run, 2)
    run.font.color.rgb = RGBColor(51, 102, 153)

    if hasattr(subsection, 'content') and subsection.content.status == 'generated':
        _add_subsection_content(doc, subsection.content.content_json)

    for topic_num, topic in enumerate(subsection.topics.filter(is_approved=True), start=1):
        progress_callback(
            f"Writing topic {section_num}.{subsection_num}.{topic_num}: {topic.title}"
        )
        _add_topic(doc, topic, section_num, subsection_num, topic_num, progress_callback)


def _add_subsection_content(doc, content_json):
    """Add subsection content"""

    if 'subsection_introduction' in content_json:
        for text in content_json['subsection_introduction'].get('paragraphs', []):
            para = doc.add_paragraph(text)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _add_topic(doc, topic, section_num, subsection_num, topic_num, progress_callback):
    """Add topic"""

    heading = doc.add_heading(f"{section_num}.{subsection_num}.{topic_num} {topic.title}", level=3)
    run = heading.runs[0]

    style_heading(run, 3)
    run.font.color.rgb = RGBColor(102, 153, 204)

    if hasattr(topic, 'content') and topic.content.status == 'generated':
        _add_topic_content(doc, topic.content.content_json, progress_callback)


def _add_topic_content(doc, content_json, progress_callback):
    """Add topic content"""

    for section in content_json.get('sections', []):
        for block in section.get('content_blocks', []):
            _add_content_block(doc, block, progress_callback)


def _add_content_block(doc, block, progress_callback):
    """Add content block"""

    # Handle different content block types
    block_type = block.get('type')

    if block_type == 'paragraph':
        para = doc.add_paragraph(block.get('content', ''))
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    elif block_type == 'bullet_list':
        for item in block.get('content', []):
            doc.add_paragraph(item, style='List Bullet')

    elif block_type == 'visual_placeholder':
        _add_visual(doc, block, progress_callback)


def _add_visual(doc, block, progress_callback):
    """Add visual"""

    # Render plotly figures or tables
    visual = block.get("generated_visual", {})

    if visual.get("status") != "ok":
        return

    fig_json = visual.get("figure_json")

    if not fig_json:
        return

    if visual.get("visual_spec", {}).get("type") == "table":
        progress_callback("Rendering table visual for the document.")
        _render_table(doc, visual)
        return

    try:
        progress_callback("Rendering chart visual for the document.")
        fig = pio.from_json(fig_json)
        image_bytes = pio.to_image(fig, format="png")
        doc.add_picture(BytesIO(image_bytes), width=Inches(6))
    except Exception:
        pass


def _render_table(doc, visual):
    """Render table"""

    # Build table from plotly json
    try:
        fig = json.loads(visual.get("figure_json"))
    except Exception:
        return

    data = fig.get("data", [])
    if not data:
        return

    table_data = data[0]

    headers = table_data.get("header", {}).get("values", [])
    cells = table_data.get("cells", {}).get("values", [])

    if not headers or not cells:
        return

    table = doc.add_table(rows=len(cells[0]) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    _apply_table_borders(table)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(header)
        _format_table_header_cell(cell)

    for row_idx in range(len(cells[0])):
        for col_idx in range(len(headers)):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cells[col_idx][row_idx])
            _format_table_body_cell(cell)


def decode_bdata(bdata, dtype):
    """Decode binary data"""

    binary = base64.b64decode(bdata)

    if dtype == "f8":
        return list(struct.unpack(f"{len(binary)//8}d", binary))
    if dtype == "f4":
        return list(struct.unpack(f"{len(binary)//4}f", binary))
    if dtype == "i4":
        return list(struct.unpack(f"{len(binary)//4}i", binary))
    if dtype == "i2":
        return list(struct.unpack(f"{len(binary)//2}h", binary))
    if dtype == "i1":
        return list(struct.unpack(f"{len(binary)}b", binary))

    return []


def style_heading(run, level):
    """Style heading"""

    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True

    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True

    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True


def _apply_table_borders(table):
    """Ensure DOCX tables render with visible row and column borders."""

    table_pr = table._tbl.tblPr
    borders = table_pr.first_child_found_in("w:tblBorders")

    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)

        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def _format_table_header_cell(cell):
    """Apply header styling so table headers are clearly separated."""

    _set_cell_shading(cell, "D9E6F5")
    _set_cell_margins(cell)
    _set_cell_borders(cell)

    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10.5)


def _format_table_body_cell(cell):
    """Apply body cell styling for better readability."""

    _set_cell_margins(cell)
    _set_cell_borders(cell)

    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.size = Pt(10)


def _set_cell_shading(cell, fill):
    """Apply background shading to a DOCX table cell."""

    cell_pr = cell._tc.get_or_add_tcPr()
    shading = cell_pr.first_child_found_in("w:shd")

    if shading is None:
        shading = OxmlElement("w:shd")
        cell_pr.append(shading)

    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    """Add padding inside DOCX table cells."""

    cell_pr = cell._tc.get_or_add_tcPr()
    margins = cell_pr.first_child_found_in("w:tcMar")

    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_pr.append(margins)

    for side, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        margin = margins.find(qn(f"w:{side}"))
        if margin is None:
            margin = OxmlElement(f"w:{side}")
            margins.append(margin)

        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def _set_cell_borders(cell):
    """Apply explicit borders to each DOCX table cell for consistent rendering."""

    cell_pr = cell._tc.get_or_add_tcPr()
    borders = cell_pr.first_child_found_in("w:tcBorders")

    if borders is None:
        borders = OxmlElement("w:tcBorders")
        cell_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)

        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "10")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
