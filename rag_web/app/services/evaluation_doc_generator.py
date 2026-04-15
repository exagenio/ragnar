from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO


def generate_evaluation_document(report, topic_evals):

    doc = Document()

    # =========================
    # TITLE
    # =========================
    title = doc.add_heading("Evaluation Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Report ID: {report.id}")
    doc.add_paragraph(f"Report Title: {report.title}")
    doc.add_paragraph()

    # =========================
    # PROJECT LEVEL CALCULATION
    # =========================
    overall_with_conc = []
    overall_without_conc = []

    agg_correctness = []
    agg_relevance = []
    agg_hallucination = []
    agg_conciseness = []

    # 🔥 keep per-topic computed rows for the table
    topic_rows = []

    for eval_obj in topic_evals:
        scores = eval_obj.geval_scores or {}

        c = float(scores.get("correctness", 0))
        r = float(scores.get("relevance", 0))
        h = float(scores.get("hallucination", 0))
        con = float(scores.get("conciseness", 0))

        overall1 = (c + r + h + con) / 4
        overall2 = (c + r + h) / 3

        overall_with_conc.append(overall1)
        overall_without_conc.append(overall2)

        agg_correctness.append(c)
        agg_relevance.append(r)
        agg_hallucination.append(h)
        agg_conciseness.append(con)

        # 🔥 store for table
        topic_rows.append({
            "id": eval_obj.topic.id,
            "title": eval_obj.topic.title,
            "overall_with": round(overall1, 2),
            "overall_without": round(overall2, 2),
            "hallucination": h,
            "correctness": c,
            "relevance": r,
            "conciseness": con,
        })

    def avg(arr):
        return round(sum(arr) / len(arr), 2) if arr else 0

    # =========================
    # SUMMARY
    # =========================
    doc.add_heading("Overall Evaluation Summary", level=1)

    doc.add_paragraph(f"Overall Score (with conciseness): {avg(overall_with_conc)}")
    doc.add_paragraph(f"Overall Score (without conciseness): {avg(overall_without_conc)}")

    doc.add_paragraph(f"Correctness: {avg(agg_correctness)}")
    doc.add_paragraph(f"Relevance: {avg(agg_relevance)}")
    doc.add_paragraph(f"Hallucination: {avg(agg_hallucination)}")
    doc.add_paragraph(f"Conciseness: {avg(agg_conciseness)}")

    # =========================
    # 🔥 TABLE SECTION (NEW)
    # =========================
    doc.add_page_break()
    doc.add_heading("Topic Evaluation Summary Table", level=1)

    headers = [
        "Topic ID",
        "Topic",
        "Overall (With Conciseness)",
        "Overall (Without Conciseness)",
        "Hallucination",
        "Correctness",
        "Relevance",
        "Conciseness",
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    # Data rows
    for row in topic_rows:
        cells = table.add_row().cells

        values = [
            row["id"],
            row["title"],
            row["overall_with"],
            row["overall_without"],
            row["hallucination"],
            row["correctness"],
            row["relevance"],
            row["conciseness"],
        ]

        for i, val in enumerate(values):
            cell = cells[i]

            if isinstance(val, float):
                cell.text = f"{val:.2f}"
            else:
                cell.text = str(val)

            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_page_break()

    # =========================
    # TOPIC DETAILS
    # =========================
    doc.add_heading("Topic-Level Evaluation", level=1)

    for eval_obj in topic_evals:
        topic = eval_obj.topic
        scores = eval_obj.geval_scores or {}

        doc.add_heading(f"Topic {topic.id}: {topic.title}", level=2)

        doc.add_paragraph("Scores:")
        doc.add_paragraph(f"Correctness: {scores.get('correctness', 0)}")
        doc.add_paragraph(f"Relevance: {scores.get('relevance', 0)}")
        doc.add_paragraph(f"Hallucination: {scores.get('hallucination', 0)}")
        doc.add_paragraph(f"Conciseness: {scores.get('conciseness', 0)}")

        doc.add_paragraph("Issues:")
        for issue in eval_obj.geval_issues or []:
            doc.add_paragraph(f"- {issue}")

        doc.add_paragraph("Summary:")
        doc.add_paragraph(eval_obj.geval_summary or "-")

        doc.add_paragraph()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer